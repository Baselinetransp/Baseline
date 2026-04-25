#!/usr/bin/env python3
"""
Convert DOCX file to Markdown format
"""

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
import sys
import os

def convert_docx_to_markdown(docx_path, output_path):
    """
    Convert a DOCX file to Markdown format
    """
    try:
        # Load the document
        doc = Document(docx_path)

        markdown_lines = []

        # Process each paragraph
        for para in doc.paragraphs:
            text = para.text.strip()

            if not text:
                # Empty paragraph - add blank line
                markdown_lines.append("")
                continue

            # Check paragraph style for headings
            style_name = para.style.name.lower()

            if 'heading 1' in style_name:
                markdown_lines.append(f"# {text}")
            elif 'heading 2' in style_name:
                markdown_lines.append(f"## {text}")
            elif 'heading 3' in style_name:
                markdown_lines.append(f"### {text}")
            elif 'heading 4' in style_name:
                markdown_lines.append(f"#### {text}")
            elif 'heading 5' in style_name:
                markdown_lines.append(f"##### {text}")
            elif 'heading 6' in style_name:
                markdown_lines.append(f"###### {text}")
            else:
                # Process inline formatting
                formatted_text = ""
                for run in para.runs:
                    run_text = run.text

                    # Apply bold
                    if run.bold:
                        run_text = f"**{run_text}**"

                    # Apply italic
                    if run.italic:
                        run_text = f"*{run_text}*"

                    formatted_text += run_text

                # Check if it's a list item (simple heuristic)
                if formatted_text.strip().startswith('•') or formatted_text.strip().startswith('-'):
                    markdown_lines.append(formatted_text)
                else:
                    markdown_lines.append(formatted_text if formatted_text else text)

            markdown_lines.append("")  # Add blank line after each paragraph

        # Process tables
        if doc.tables:
            for table in doc.tables:
                markdown_lines.append("")
                # Process table headers
                if table.rows:
                    header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                    markdown_lines.append("| " + " | ".join(header_cells) + " |")
                    markdown_lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")

                    # Process table rows
                    for row in table.rows[1:]:
                        cells = [cell.text.strip() for cell in row.cells]
                        markdown_lines.append("| " + " | ".join(cells) + " |")

                markdown_lines.append("")

        # Write to output file
        markdown_content = "\n".join(markdown_lines)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)

        print(f"Successfully converted '{docx_path}' to '{output_path}'")
        print(f"Output file size: {len(markdown_content)} characters")

        return True

    except Exception as e:
        print(f"Error: Error converting document: {e}", file=sys.stderr)
        return False

if __name__ == "__main__":
    # Input and output paths
    input_file = "Transport_Hub_Requirements_Engineering_Document.docx"
    output_file = "Transport_Hub_Requirements_Engineering_Document.md"

    # Convert
    if os.path.exists(input_file):
        convert_docx_to_markdown(input_file, output_file)
    else:
        print(f"Error: Error: File '{input_file}' not found", file=sys.stderr)
        sys.exit(1)
